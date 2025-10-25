# app.py
import os
import re
import time
import json
import shutil
import tempfile
import unicodedata
from io import BytesIO

import streamlit as st
import pandas as pd
import requests
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

import yt_dlp
from moviepy.editor import VideoFileClip
from tqdm import tqdm
import zipfile

# ------------------- Utilities -------------------

st.set_page_config(page_title="SocialScribe (Streamlit)", layout="wide")

def sanitize_name(name):
    if not name:
        return "Unknown"
    name = unicodedata.normalize("NFKC", str(name))
    name = re.sub(r'[\\/*?:"<>|]', "", name)
    name = name.rstrip(". ")
    name = "".join(ch for ch in name if ch.isprintable())
    return name.strip() or "Unknown"

def set_kannada_font(run):
    try:
        run.font.name = "Nirmala UI"
        rPr = run._element.rPr
        rFonts = rPr.rFonts
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Nirmala UI')
    except Exception:
        pass

def set_cell_border(cell, **kwargs):
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                border_element = tcBorders.find(qn(tag))
                if border_element is None:
                    border_element = OxmlElement(tag)
                    tcBorders.append(border_element)
                for key, val in edge_data.items():
                    border_element.set(qn(f'w:{key}'), str(val))
    except Exception:
        pass

def add_metadata_to_cell(cell, label, value, is_link=False):
    p = cell.add_paragraph()
    p_format = p.paragraph_format
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(6)

    run = p.add_run(label + " ")
    run.bold = True
    try:
        run.font.color.rgb = RGBColor(255, 0, 0)
    except Exception:
        pass
    set_kannada_font(run)

    run_val = p.add_run(value)
    if is_link:
        try:
            run_val.font.color.rgb = RGBColor(5, 99, 193)
            run_val.font.underline = True
        except Exception:
            pass
    set_kannada_font(run_val)
    return p

# ------------------- YouTube report generator -------------------

def fetch_ytdl_info(url, cookies_file=None):
    meta_opts = {"quiet": True, "skip_download": True, "noplaylist": True}
    if cookies_file:
        meta_opts["cookiefile"] = cookies_file
    with yt_dlp.YoutubeDL(meta_opts) as ydl:
        info = ydl.extract_info(url, download=False)
    return info

def download_thumbnail_from_info(info, out_path):
    # Try ytdl thumbnail first
    thumb = info.get("thumbnail") or (info.get("thumbnails") and info.get("thumbnails")[-1].get("url"))
    if not thumb:
        return False
    try:
        r = requests.get(thumb, timeout=15)
        r.raise_for_status()
        img = Image.open(BytesIO(r.content)).convert("RGB")
        img.save(out_path)
        return True
    except Exception:
        return False

def build_report_from_urls(urls):
    doc = Document()
    # margins similar to original
    for section in doc.sections:
        try:
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        except Exception:
            pass

    tmpdir = tempfile.mkdtemp(prefix="socialscribe_report_")
    saved_images = []

    for url in urls:
        try:
            info = fetch_ytdl_info(url)
        except Exception as e:
            st.warning(f"Could not fetch metadata for: {url} | {e}")
            continue

        title = info.get("title", "Unknown Title")
        uploader = info.get("uploader") or info.get("channel") or "Unknown Channel"
        upload_date_raw = info.get("upload_date") or info.get("upload_date")
        upload_date = ""
        if upload_date_raw and isinstance(upload_date_raw, str) and len(upload_date_raw) >= 8:
            # format YYYYMMDD
            try:
                upload_date = f"{upload_date_raw[6:8]}-{upload_date_raw[4:6]}-{upload_date_raw[0:4]}"
            except Exception:
                upload_date = upload_date_raw

        # metadata table
        doc.add_paragraph()
        meta_table = doc.add_table(rows=1, cols=1)
        try:
            meta_table.columns[0].width = Inches(7.7)
        except Exception:
            pass
        cell = meta_table.cell(0, 0)
        cell.text = ""
        add_metadata_to_cell(cell, "Title:", title)
        add_metadata_to_cell(cell, "Post Date:", upload_date)
        add_metadata_to_cell(cell, "Link:", url, is_link=True)
        add_metadata_to_cell(cell, "Channel:", uploader, is_link=False)
        set_cell_border(cell, top={"sz":12,"val":"single","color":"#000000"},
                             bottom={"sz":12,"val":"single","color":"#000000"},
                             left={"sz":12,"val":"single","color":"#000000"},
                             right={"sz":12,"val":"single","color":"#000000"})

        # thumbnail
        thumb_path = os.path.join(tmpdir, sanitize_name(title)[:100] + "_thumb.jpg")
        ok = download_thumbnail_from_info(info, thumb_path)
        if ok and os.path.exists(thumb_path):
            p = doc.add_paragraph()
            try:
                p.add_run().add_picture(thumb_path, height=Inches(2.5), width=Inches(3.3))
            except Exception:
                try:
                    p.add_run().add_picture(thumb_path, height=Inches(2.0))
                except Exception:
                    pass
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            saved_images.append(thumb_path)

    # save doc to bytes
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    # cleanup
    try:
        shutil.rmtree(tmpdir)
    except Exception:
        pass

    return bio

# ------------------- Trimmer (Excel -> clips -> zip) -------------------

def build_ydl_opts(output_dir, cookies_file=None):
    opts = {
        "format": "best[height<=720]/best",
        "outtmpl": os.path.join(output_dir, "temp_video.%(ext)s"),
        "quiet": True,
        "noplaylist": True,
    }
    if cookies_file:
        opts["cookiefile"] = cookies_file
    return opts

def download_video(url, output_dir, cookies_file=None):
    opts = build_ydl_opts(output_dir, cookies_file)
    with yt_dlp.YoutubeDL(opts) as ydl:
        info = ydl.extract_info(url, download=True)
        path = ydl.prepare_filename(info)
    return info, path

def parse_timecode(tc):
    tc = str(tc).strip().replace(".", ":")
    parts = [int(p) for p in tc.split(":")]
    if len(parts) == 1:
        return parts[0]
    if len(parts) == 2:
        return parts[0]*60 + parts[1]
    if len(parts) == 3:
        return parts[0]*3600 + parts[1]*60 + parts[2]
    return None

def trim_clip(input_file, start_s, end_s, output_file):
    with VideoFileClip(input_file) as video:
        start_s = max(0, float(start_s))
        end_s = min(float(video.duration), float(end_s))
        if start_s >= end_s:
            return False
        clip = video.subclip(start_s, end_s)
        # force write to disk (moviepy doesn't stream to BytesIO reliably)
        clip.write_videofile(output_file, codec="libx264", audio_codec="aac", logger=None)
        clip.close()
    return True

def process_excel_to_zip(excel_bytes, cookies_file=None, max_workers=2, sleep_between_clips=0.2):
    # excel_bytes can be BytesIO or file path; here we accept bytes
    tmp_root = tempfile.mkdtemp(prefix="socialscribe_trim_")
    out_tmp = os.path.join(tmp_root, "outputs")
    os.makedirs(out_tmp, exist_ok=True)

    # read excel
    df = pd.read_excel(excel_bytes)
    df = df.dropna(how='all')
    columns = df.columns.tolist()

    created_files = []

    for idx, row in df.iterrows():
        source_identifier = str(row[columns[0]]).strip()
        if not source_identifier:
            continue
        is_local_file = os.path.exists(source_identifier)
        video_path = None
        title = "Untitled"
        channel = "Unknown"

        if is_local_file:
            video_path = source_identifier
            title = os.path.splitext(os.path.basename(video_path))[0]
            custom_channel = row.get(columns[1]) if len(columns) > 1 else None
            if custom_channel and not pd.isna(custom_channel) and str(custom_channel).strip():
                channel = str(custom_channel).strip()
            else:
                channel = os.path.basename(os.path.dirname(video_path)) or "LocalFile"
        else:
            # try to fetch metadata then download
            try:
                meta_opts = {"quiet": True, "skip_download": True, "noplaylist": True}
                if cookies_file:
                    meta_opts["cookiefile"] = cookies_file
                with yt_dlp.YoutubeDL(meta_opts) as ydl:
                    info = ydl.extract_info(source_identifier, download=False)
                    title = info.get("title", title)
                    channel = info.get("uploader") or info.get("channel") or channel
            except Exception as e:
                st.warning(f"Could not fetch metadata for {source_identifier}: {e}")
                continue

        clean_channel = sanitize_name(channel)
        clean_title = sanitize_name(title)
        folder_name = f"{clean_channel}_{clean_title}"[:180]
        output_folder = os.path.join(out_tmp, folder_name)
        os.makedirs(output_folder, exist_ok=True)

        if not is_local_file:
            try:
                _, video_path = download_video(source_identifier, output_folder, cookies_file=cookies_file)
            except Exception as e:
                st.warning(f"Failed to download {source_identifier}: {e}")
                continue

        # columns with timecodes start from 2 if local file else 1 (like original app)
        start_col_index = 2 if is_local_file else 1
        for col in columns[start_col_index:]:
            cell = row.get(col)
            if pd.isna(cell):
                continue
            text = str(cell).strip()
            if not text or "-" not in text:
                continue
            parts = re.split(r'[/,]', text)
            for idxp, part in enumerate(parts):
                if "-" not in part:
                    continue
                start_str, end_str = part.split("-", 1)
                ssec = parse_timecode(start_str)
                esec = parse_timecode(end_str)
                if ssec is None or esec is None:
                    continue
                col_clean = sanitize_name(col.split("(")[0].strip())
                filename = f"{col_clean}_{clean_channel}_{clean_title}_part{idxp+1}.mp4"
                outpath = os.path.join(output_folder, filename)
                if not os.path.exists(outpath):
                    try:
                        trim_clip(video_path, ssec, esec, outpath)
                        created_files.append(outpath)
                    except Exception as e:
                        st.warning(f"Error trimming {outpath}: {e}")
                time.sleep(sleep_between_clips)

        # cleanup downloaded video
        try:
            if not is_local_file and video_path and os.path.exists(video_path):
                os.remove(video_path)
        except Exception:
            pass

    # create zip from created_files
    zip_bio = BytesIO()
    with zipfile.ZipFile(zip_bio, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for fpath in created_files:
            arcname = os.path.relpath(fpath, out_tmp)
            zf.write(fpath, arcname)
    zip_bio.seek(0)

    # cleanup tempdir
    try:
        shutil.rmtree(tmp_root)
    except Exception:
        pass

    return zip_bio

# ------------------- Streamlit UI -------------------

st.markdown("<h1 style='color:#1e73d1'>SocialScribe</h1>", unsafe_allow_html=True)
st.write("YouTube Report Generator & Video Trimmer — web edition")

tabs = st.tabs(["Report Generator (URLs)", "Video Trimmer (Excel → ZIP)"])

# -------- Report Generator Tab --------
with tabs[0]:
    st.subheader("Generate Word report from YouTube URLs")
    st.info("Paste one YouTube URL per line. Thumbnails are downloaded via yt-dlp where possible.")
    urls_text = st.text_area("Paste YouTube URLs (one per line)", height=180)
    urls = [u.strip() for u in urls_text.splitlines() if u.strip()]

    col1, col2 = st.columns([1, 4])
    with col1:
        gen_btn = st.button("Generate .docx Report")
    with col2:
        st.write("When ready, click Generate. The Word file will be available for download.")

    if gen_btn:
        if not urls:
            st.warning("Please paste at least one YouTube URL.")
        else:
            with st.spinner("Building report..."):
                doc_bytes = build_report_from_urls(urls)
                st.success("Report ready.")
                st.download_button("Download Word Report (.docx)", data=doc_bytes.getvalue(),
                                   file_name="SocialScribe_YouTube_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# -------- Trimmer Tab --------
with tabs[1]:
    st.subheader("Process Excel to trim clips and download ZIP")
    st.info("Upload Excel where first column is YouTube URL or local path. Timecodes as columns (e.g. 00:01-00:12) — multiple ranges separated by comma or slash.")
    uploaded = st.file_uploader("Upload Excel (.xlsx)", type=["xlsx"])
    run_btn = st.button("Process and Create ZIP")
    if uploaded:
        st.write("File ready:", uploaded.name)
    if run_btn:
        if not uploaded:
            st.warning("Upload an Excel file first.")
        else:
            with st.spinner("Processing Excel and trimming clips (this may take time)..."):
                try:
                    # stream uploaded file into BytesIO
                    excel_bytes = uploaded.read()
                    zip_bio = process_excel_to_zip(BytesIO(excel_bytes))
                    st.success("ZIP ready.")
                    st.download_button("Download Clips ZIP", data=zip_bio.getvalue(), file_name="SocialScribe_clips.zip", mime="application/zip")
                except Exception as e:
                    st.error(f"Processing failed: {e}")

st.markdown("---")
st.write("Notes: heavy video work may be slow on free hosting. For large batch jobs consider running locally or on a machine with more CPU / disk.")
